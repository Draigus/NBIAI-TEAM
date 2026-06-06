"""
Build a Word document from the CH Zone Architecture Design Spec and Backend Architecture Diagram.
Combines both into a single professional .docx with Word-native formatting.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x2d, 0x2d, 0x3f)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)

ACCENT = RGBColor(0x6c, 0x5c, 0xe7)
DARK = RGBColor(0x2d, 0x2d, 0x3f)
MUTED = RGBColor(0x66, 0x66, 0x80)
RED = RGBColor(0xcc, 0x33, 0x33)
GREEN = RGBColor(0x00, 0x88, 0x80)
BLUE = RGBColor(0x33, 0x66, 0xaa)

def add_para(text, bold=False, italic=False, color=None, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size: run.font.size = size
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_callout(title, body, border_color='4472C4'):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    run_title = p.add_run(title + '\n')
    run_title.bold = True
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(*(int(border_color[i:i+2], 16) for i in (0, 2, 4)))
    run_body = p.add_run(body)
    run_body.font.size = Pt(10)
    return p

def shade_cell(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        shade_cell(cell, '2D2D3F')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri % 2 == 1:
                shade_cell(cell, 'F2F2F7')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    return p

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('COUCH HEROES', bold=True, color=ACCENT, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Zone Population Architecture', bold=True, color=DARK, size=Pt(28), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Design Specification + Backend Architecture', color=MUTED, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('Prepared by NBI Gaming, Glen Pryer, CPO (Advisory)', color=MUTED, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Version 3.0  /  June 2026', color=MUTED, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('Audience: Vardis (CEO), Aris (COO), Mustafa (Head of Tech), Robin (Game Director)', color=MUTED, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Internal / Confidential', bold=True, color=RED, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 0. DECISION FRAMEWORK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('0. Decision Framework', level=1)
doc.add_paragraph('We need to lock in a zone population target so the tech team can start prototyping the networking layer. How many concurrent players per zone are we building for?')
doc.add_paragraph('It is not a single number. It is a spectrum, and each tier carries a different engineering investment, a different team composition, and a different competitive position.')

add_table(
    ['Target', 'Engineering Effort', 'Team Requirement', 'Competitive Position', 'Risk'],
    [
        ['100 to 150\nper zone',
         'Significant. Replication Graph plugin with custom C++ node implementation, spatial partitioning, dormancy config, per-actor frequency tuning.',
         '1 to 2 engineers with UE5 Replication Graph experience. Custom C++, not configuration.',
         'Parity with Fortnite (100 players, built by Epic\'s own engine team). Below FFXIV (~300).',
         'Moderate'],
        ['200 to 300\nper zone\n[RECOMMENDED]',
         'Major. Everything above plus custom interest management tiers, hub sharding, staggered replication, state transfer pipeline. Multi-quarter infrastructure build.',
         '2 to 3 senior engineers with UE5 dedicated server experience, dedicated to networking.',
         'Approaches FFXIV\'s commonly observed ~300 zone cap. Would be rare among shipped Social MMOs. Unusual for a UE5 title with real-time combat.',
         'Moderate\nto High'],
    ]
)

add_callout('RECOMMENDATION',
    'Build for 200 to 300 per zone. Prototype at 200 first. The prototype tells us where the real ceiling is before we commit production time.',
    '6C5CE7')

add_callout('ONE THING EVERYONE NEEDS TO BE CLEAR ON',
    'Zone population and combat population are different numbers. Any single combat encounter (world boss, corruption event, PvP fight) caps at about 100 participants. Players beyond that observe but cannot join the fight. The zone holds 200 to 300 players coexisting across exploration, social, and combat. The combat happens in clusters of up to 100 within that zone. This is the standard MMO pattern and it is how we should talk about it externally.',
    'CC3333')

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph('Our design pillars (discovery, corruption, no-stat cosmetics, soft-class, social-first) all depend on the same thing: the world needs to feel alive and full of other players. The GDD originally specced 30 to 40 per instance. That is not enough for a Social MMO. This document lays out what it takes to get to 200 to 300 per zone, with Downtime sharding dynamically as it fills up, and combat encounters running smoothly at up to 100 participants.')

doc.add_heading('Why This Matters', level=2)
doc.add_paragraph('The social experience is the product. A world with 30 players per zone feels empty. Palia is the cautionary example: its low population cap was frequently cited by players as harming the social feel, and the studio\'s commercial outcomes speak for themselves. Our five factions only feel real when faction districts have people in them. Corruption only generates shared stories when multiple players see it happen. The cosmetic system only drives aspiration when players see each other\'s gear. All of that requires density.')

doc.add_paragraph('Social MMO competitors with gameplay systems cap at 25 to 50 per zone (Palia, Tower of Fantasy, Destiny 2). Social virtual worlds without combat reach higher (VRChat at 80, Second Life at 100) but lack game systems. FFXIV is commonly observed around approximately 300 per zone (not officially published by Square Enix), with combat heavily instanced and a long-mature custom engine. A UE5 Social MMO that delivers 200 to 300 with real-time open-world combat would be approaching FFXIV\'s zone density while supporting open-world PvE and PvP on the zone server.')

# ═══════════════════════════════════════════════════════════════
# 2. COMPETITIVE LANDSCAPE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Competitive Landscape', level=1)
doc.add_paragraph('Figures below mix official documentation, developer statements, public reporting, and community observation. Each row labels the source confidence. Community-reported figures should be treated as directional, not independently verified.')

add_table(
    ['Title', 'Engine', 'Zone Model', 'Combat', 'Source'],
    [
        ['Palia', 'UE4', '25 per instance', 'No combat', 'Community reported'],
        ['Tower of Fantasy', 'UE4', '~30 to 50 per channel', 'Open-world + instanced', 'Community reported'],
        ['Throne and Liberty', 'UE4', 'Hundreds in siege (marketing)', 'Large PvP sieges', 'Marketing claim'],
        ['Destiny 2', 'Custom', '20 to 26 in social hub', 'Instanced PvE/PvP', 'Community reported'],
        ['FFXIV', 'Custom', '~300 per zone (commonly observed, not official)', 'Instanced (8 to 24 PvE, 72 PvP)', 'Community reported'],
        ['World of Warcraft', 'Custom', 'Phasing/sharding/layering, caps not published', '20-player raids, 40 BGs', 'Community reported'],
        ['Guild Wars 2', 'Custom', 'Megaserver dynamic instancing, WvW hundreds', 'Large-scale WvW, 5v5 structured', 'Dev stated'],
        ['Ashes of Creation', 'UE5', '~300 in alpha stress test (dev stated)', '100v100 siege target', 'Dev stated'],
        ['Mortal Online 2', 'UE5', '~500 seamless battle (dev claim)', 'Open-world PvP', 'Marketing claim'],
        ['VRChat', 'Unity', '80 hard cap per instance', 'No combat', 'Official docs'],
        ['Second Life', 'Custom', '100 basic + premium allowance per region', 'Minimal', 'Official docs'],
    ]
)

doc.add_heading('Key Lessons', level=2)
add_bullet('Palia\'s 25-player cap was widely criticised for making the social world feel empty. Commercial outcomes followed.', bold_prefix='Empty World Risk:')
add_bullet('Throne and Liberty marketed large sieges but community reported severe performance issues. Consolidated from 107 to 25 servers ~5 months post-launch.', bold_prefix='Performance Promise Risk:')
add_bullet('FFXIV\'s visible numbered hub instances are accepted by players at scale. This is the model for Downtime sharding.', bold_prefix='FFXIV Model:')
add_bullet('GW2\'s megaserver scoring algorithm (party, guild, language) and centralised Trading Post are the models for social routing and cross-shard commerce.', bold_prefix='GW2 Model:')

# ═══════════════════════════════════════════════════════════════
# 3. ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Architecture Overview', level=1)
doc.add_paragraph('The architecture uses five distinct server profiles, each tuned for its activity type. Different activities have fundamentally different networking requirements. Social interactions are cheap. Combat is expensive. Housing is persistent and low-density. Dungeons and PvP need guaranteed tick rates with small player counts.')

doc.add_heading('Five Server Profiles', level=2)

doc.add_heading('Profile 1: Overworld Zone Servers', level=3)
doc.add_paragraph('One UE5 dedicated server process per zone. 200 to 300 players. 10 to 20 Hz server tick rate. Replication Graph with custom spatial nodes (must be built, not stock). 3-tier interest management: Near (3 to 10 Hz effective), Mid (1 to 3 Hz), Far (0.5 to 1 Hz). Combat clusters elevated to priority 3.0. Staggered replication: 50 to 100 connections per tick, rotated. Corruption Manager actor with negligible replication work between changes, versioned snapshots for late joins. Hardware: c7i.4xlarge (16 vCPU, 32 GB).')

doc.add_heading('Profile 2: Downtime Hub Shards', level=3)
doc.add_paragraph('Approximately 200 players per shard (to be validated by prototype). Social-graph-aware routing keeps parties, guilds, friends, and faction members together. No combat, no AI, no physics. New shards spawn at capacity. Five faction districts distribute density naturally. Hardware: c7i.2xlarge (8 vCPU, 16 GB).')

doc.add_heading('Profile 3: User Space Housing Servers', level=3)
doc.add_paragraph('Shared processes hosting 20 to 50 concurrent occupied User Spaces (initial sizing hypothesis, must be validated). 12 to 24 visitors per house. UE5 World Partition streaming, only active sub-levels loaded. No combat, no AI. Static furniture is persisted placement data, not live-replicated actors. Real load is interaction state and serialisation. Hardware: c7i.xlarge (4 vCPU, 8 GB).')

doc.add_heading('Profile 4: Instance Servers', level=3)
doc.add_paragraph('Dungeons (4-player, 30 Hz), Raids (8 to 24, 30 Hz), PvP Arena (3v3 to 20v20, 30 Hz), Siege/Territory War (50v50 target, 20 Hz target). These are prototype targets, not commitments. Spawned on demand by Agones from a warm pool.')

doc.add_heading('Profile 5: Backend Services', level=3)
add_bullet('Agones on Kubernetes for server lifecycle management (open source, cloud-agnostic)')
add_bullet('Redis (ephemeral state transfer cache) + PostgreSQL (durable authority for all persistent state)')
add_bullet('Custom metagame backend (accounts, social graph, matchmaking, economy, faction state). No third-party managed service.')
add_bullet('Corruption world-state service (pub/sub to zone servers, eventual consistency)')
add_bullet('EasyAntiCheat (client-side anti-cheat, base free via EOS) + Cloudflare Spectrum Enterprise (DDoS/IP hiding)')
add_bullet('Telemetry service (per-connection bandwidth, tick time, event participation, dashboarded via Grafana)')

doc.add_heading('Player Distribution Across a 300-Player Zone Ecosystem', level=2)
doc.add_paragraph('A "300-player zone" refers to the total ecosystem of server processes serving one game zone, not a single server process. A plausible busy-state distribution:')
add_bullet('40 to 80 in Downtime areas (on hub shard servers, separate processes)')
add_bullet('30 to 80 in active combat clusters (corruption events, world bosses, PvP)')
add_bullet('120 to 200 exploring, questing, gathering, fishing, spread across the zone')
add_bullet('20 to 50 transitioning to or from instanced content')
p = doc.add_paragraph()
run = p.add_run('These are illustrative proportions, not hard allocations.')
run.italic = True
run.font.color.rgb = MUTED

# ═══════════════════════════════════════════════════════════════
# 4. CORRUPTION STATE REPLICATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Corruption Mechanic: Networking Design', level=1)
doc.add_paragraph('Corruption is one of our five design pillars. It changes zone visuals, affects accessible areas, modifies enemy spawns, and drives core quest content. For networking, corruption is world state that changes over time and must be visually consistent across all players in a zone.')

doc.add_heading('Tiered Hybrid Model', level=2)
doc.add_paragraph('Layer 1: Per-zone corruption scalar (0 to 100, one byte). The authoritative gameplay value driving spawns, accessibility gates, and VFX palette. Replicated to all clients in periodic world-state snapshots. Negligible bandwidth.')
doc.add_paragraph('Layer 2: Low-resolution spatial grid (64 x 64 cells, 1 byte per cell = 4 KB total). Visual corruption spread patterns. Replicated via delta compression (only changed cells sent). Typical update: 50 to 200 bytes every 5 to 10 seconds, usually smaller than one full moving character update.')
doc.add_paragraph('Corruption grids need versioned snapshots (so late-joining players receive the current state), idempotent cell updates (so duplicate cleansing events do not corrupt the grid), and full-state recovery for players who join mid-event.')
doc.add_paragraph('Cross-shard consistency: corruption state owned by a central world-state service backed by PostgreSQL. Changes fan out via pub/sub with 1 to 5 second latency. Acceptable for environmental state.')

# ═══════════════════════════════════════════════════════════════
# 5. USER SPACE (HOUSING)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. User Space (Housing) Architecture', level=1)
doc.add_paragraph('User Spaces are personal pocket dimensions accessed via the House of Houses. The initial sizing hypothesis is 20 to 50 concurrent occupied User Spaces per server process. This must be validated with furniture count, visitor count, memory, serialisation, and persistence tests.')
doc.add_paragraph('Visitor cap: 12 to 24 per User Space (consistent with ESO and FFXIV). Static furniture should be persisted placement data loaded on entry, not live-replicated actors. The actual server load comes from interaction state, serialisation of placement data on entry, and load time for complex User Spaces.')
doc.add_paragraph('Entry is a server transition via ClientTravel() masked by a portal animation. Expected transition time: 3 to 5 seconds.')

# ═══════════════════════════════════════════════════════════════
# 6. COMBAT NETWORKING
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Combat Networking Design', level=1)

doc.add_heading('The 100-Participant Event Cap', level=2)
doc.add_paragraph('Open-world events cap at 100 active participants. This is a designed constraint, not a hidden limitation:')
add_bullet('First 100 players who engage become active participants with elevated network priority')
add_bullet('Additional arrivals become observers (reduced visual fidelity, cannot deal damage or receive loot)')
add_bullet('When a participant leaves, the longest-waiting observer is promoted')
add_bullet('Robin\'s team should frame this as "the corruption is overwhelming, reinforcements standing by"')

doc.add_heading('Soft-Lock Hybrid Combat Model', level=2)
doc.add_paragraph('Soft-lock was selected because it provides the best balance between networking efficiency and skill expression. Server validation is simple and binary: "Was the target within lock-on range at the server timestamp?" No projectile simulation, no frame-precise hitboxes, no favour-the-shooter lag compensation needed.')

doc.add_heading('Skill Expression Techniques', level=3)
add_bullet('Flanking/rear attacks deal bonus damage (FFXIV model)', bold_prefix='Positional bonuses:')
add_bullet('Invincibility frames with precise timing windows (GW2 dodge system)', bold_prefix='Dodge i-frames:')
add_bullet('Upper body locked to target, lower body free to move and dodge', bold_prefix='Split-body animation:')
add_bullet('Lock weakens or breaks when target dashes or enters cover', bold_prefix='Lock degradation:')
add_bullet('Ability chains that reward correct order with bonus effects', bold_prefix='Combo sequencing:')
add_bullet('Cancel recovery frames into startup of next ability for skill ceiling', bold_prefix='Animation cancelling:')

doc.add_heading('Networking-Friendly Design Rules', level=3)
add_bullet('AoE abilities cap at 12 targets (reduces combat fan-out, helps avoid O(N squared) worst cases)')
add_bullet('VFX tiers per ability: full (under 30 participants), reduced (30 to 60), minimal (60+)')
add_bullet('Death respawn at safe locations away from combat sites (prevents density buildup)')
add_bullet('PvP requires opt-in flag (single replicated boolean, negligible bandwidth)')
add_bullet('Combat lockout at shard boundaries (prevents flee-to-invulnerability exploit)')

doc.add_heading('Latency Targets (prototype targets, not commitments)', level=2)
add_table(
    ['Activity', 'Max RTT Target', 'End-to-End Budget', 'Notes'],
    [
        ['Combat ability', '100ms', '150 to 200ms', 'Client prediction makes it feel instant'],
        ['Soft-lock acquisition', '150ms', '200 to 250ms', 'Binary server validation'],
        ['Movement (near tier)', '100ms', '150 to 200ms', 'Prediction + reconciliation'],
        ['Social interactions', '250ms', '300 to 500ms', 'Not latency-sensitive'],
        ['Hub shard transition', 'N/A', '3 to 5 seconds', 'Masked by animation'],
        ['Corruption update', 'N/A', '5 to 10 seconds', 'Eventual consistency'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 7. RECOMMENDED TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Recommended Technology Stack', level=1)
doc.add_paragraph('A single recommended stack. This is a starting point, not a menu.')

add_table(
    ['Layer', 'Recommendation', 'Rationale'],
    [
        ['Replication', 'Replication Graph', 'Stable, proven. GridSpatialization2D for spatial culling (stock). Distance-based frequency tiers require a custom node (must be built). Iris (Experimental in UE 5.7) should not be planned around.'],
        ['Orchestration', 'Agones on Kubernetes', 'Open source, cloud-agnostic, established game server orchestration. No vendor lock-in.'],
        ['Hosting', 'AWS c7i (x86-64)', 'UE5 servers are x86 builds. c7i.4xlarge for zones, c7i.xlarge for instances.'],
        ['Metagame', 'Custom-built backend', 'Accounts, social graph, matchmaking, economy, faction state. No vendor dependency.'],
        ['Anti-Cheat', 'EasyAntiCheat', 'Base free via EOS. Client-side integrity checks + server authority. Enterprise support may need separate agreement.'],
        ['DDoS', 'Cloudflare Spectrum Enterprise + AWS Shield Standard', 'Spectrum hides server IPs (Enterprise quote required). Shield Standard is baseline. Spectrum drops fragmented UDP: keep datagrams below path MTU.'],
    ]
)

doc.add_heading('Fallback Positions', level=2)
add_table(
    ['Layer', 'Fallback', 'When to Switch'],
    [
        ['Orchestration', 'Redwood MMO Framework', 'If Agones alone insufficient for zone routing/sharding'],
        ['Server Meshing', 'OmniMesh (StarVault, licensed)', 'Only if single-server ceiling below required target range'],
        ['Backend', 'Pragma (managed)', 'Only if custom backend proves too large a build. Adds vendor dependency.'],
    ]
)

add_callout('TECHNOLOGIES NOT RECOMMENDED',
    'Custom server meshing: multi-year investment, not viable for a 55-person studio. Third-party distributed simulation vendors are not recommended as launch dependencies without a dedicated commercial and technical due-diligence pass.',
    'CC3333')

# ═══════════════════════════════════════════════════════════════
# 8. SECURITY AND RESILIENCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Security and Resilience', level=1)

doc.add_heading('DDoS Mitigation', level=2)
doc.add_paragraph('Cloudflare Spectrum Enterprise proxies all game traffic through its global edge network, hiding game server origin IPs from clients. AWS Shield Standard does NOT hide IPs on its own. With the Agones/EKS architecture, Spectrum is required for IP hiding. Protocol design must keep UDP datagrams below path MTU because Spectrum drops fragmented UDP packets at the edge.')

doc.add_heading('Anti-Cheat', level=2)
doc.add_paragraph('Two complementary layers. Server authority validates all state-changing inputs. Interest management reduces remote-state exposure for irrelevant actors but does not replace visibility checks for relevant actors. EasyAntiCheat provides client-side anti-cheat with integrity checks and telemetry that helps detect common aimbots, tampering, and memory manipulation.')

doc.add_heading('Crash Recovery', level=2)
doc.add_paragraph('PostgreSQL is the durable authority for all important player state. Writes must be idempotent, queued/backpressured where appropriate, and load-tested against peak event cadence. Redis holds hot transient state (position, health, combat state) with 30 to 60 second snapshots for fast recovery. Redis is a cache, not the recovery authority.')
doc.add_paragraph('On crash: Agones detects health check failure (5 to 10 seconds), allocates warm-pool replacement (1 to 5 seconds), loads zone level (10 to 30 seconds), loads state from PostgreSQL + Redis. Target recovery: 60 to 120 seconds. Hot standby is excluded from baseline because it materially increases infrastructure cost and operational complexity.')

doc.add_heading('Console Compatibility', level=2)
doc.add_paragraph('Dedicated servers avoid peer-host NAT traversal, but PSN/XBL session integration, invite handling, reconnect behaviour, and TRC/XR certification requirements must be validated separately. Build platform-specific requirements in from the start.')

# ═══════════════════════════════════════════════════════════════
# 9. COST MODEL
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Comprehensive Cost Model', level=1)
doc.add_paragraph('Baseline infrastructure estimate, excluding Cloudflare Spectrum Enterprise quote and commercial support agreements. 10-zone baseline (10 zones x 300 = ~3,000 in overworld zones). For 30,000 peak CCU, zone compute and zone egress scale roughly linearly. Backend, telemetry, Redis/PostgreSQL, support, and DDoS costs require a separate capacity model.')

add_table(
    ['Component', 'Specification', 'Monthly (AWS reserved, us-east-1)'],
    [
        ['Zone servers (10)', 'c7i.4xlarge', '$3,200 to $5,200'],
        ['Hub shards (20)', 'c7i.2xlarge', '$3,200 to $5,200'],
        ['Instance servers (50)', 'c7i.xlarge warm pool', '$5,000 to $9,000'],
        ['Housing servers (10)', 'c7i.xlarge', '$1,000 to $1,800'],
        ['PostgreSQL', 'RDS db.r6g.2xlarge Multi-AZ', '$1,400 to $1,800'],
        ['Redis', 'ElastiCache 3-node cluster', '$600 to $900'],
        ['Metagame backend', '3x c7i.xlarge', '$1,500 to $2,500'],
        ['Network egress (zone traffic)', '150 to 270 TB/month', '$11,000 to $18,000'],
        ['DDoS', 'Spectrum Enterprise + Shield', 'TBD Enterprise quote'],
        ['Monitoring', 'CloudWatch, Grafana', '$500 to $1,500'],
        ['CDN', 'CloudFront', '$500 to $2,000'],
        ['CI/CD + stress testing', 'Build servers, sim clients', '$1,000 to $3,000'],
        ['Dev + staging', '2x reduced-scale', '$3,000 to $5,000'],
        ['TOTAL (excl. Spectrum)', '', '$31,000 to $56,000'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Assumptions: AWS us-east-1, 1-year reserved instances (no upfront), public internet egress at standard tiered pricing, 8h peak + 16h at 20% utilisation. Savings Plans or committed-use agreements would reduce compute. Private interconnect would reduce egress.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = MUTED

# ═══════════════════════════════════════════════════════════════
# 10. RISK REGISTER
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Risk Register', level=1)

add_table(
    ['Risk', 'Severity', 'Mitigation', 'Owner'],
    [
        ['Zone server cannot sustain 15 Hz at 300', 'HIGH', 'Prototype at 200 first. If below 300, reduce or evaluate OmniMesh.', 'Mustafa'],
        ['Launch population too low to fill zones', 'HIGH', 'MassAI NPC crowds expected to help (validate in testing). Concentrated POIs. Megaserver zone merging.', 'Robin + Mustafa'],
        ['Hub sharding breaks social connections', 'HIGH', 'Social-graph routing as primary criterion. Manual shard-switch UI.', 'Mustafa'],
        ['State transfer loses player data', 'HIGH', 'PostgreSQL durable authority. Redis 120s TTL. Verification handshake. Return to source on failure.', 'Mustafa'],
        ['Zone server crash at 300+ players', 'HIGH', 'PostgreSQL + Redis snapshots. Warm pool. 60 to 120s recovery.', 'Mustafa + DevOps'],
        ['Team lacks UE5 networking specialists', 'HIGH', 'Assess team. Hire or contract if < 2 with DS experience.', 'Vardis + Mustafa'],
        ['Custom backend exceeds capacity', 'HIGH', 'Restrict launch to core features. Defer chat/economy automation to post-launch. Load-test.', 'Backend Lead'],
        ['UE5 engine version lock-in', 'MEDIUM', 'Lock version. Migrate only between milestones.', 'Mustafa'],
        ['Reliable RPC bunch contention', 'MEDIUM', 'Code review gate: no reliable RPCs in combat paths.', 'Mustafa'],
        ['Client rendering fails at 80+ chars', 'MEDIUM', 'Impostor billboards + animation LOD on min-spec early.', 'Art Director'],
        ['Shard transitions exceed 5 seconds', 'MEDIUM', 'Prototype early. Fallback: FFXIV visible-instance model.', 'Mustafa'],
        ['DDoS on game servers', 'Inherent: HIGH\nResidual: LOW', 'Spectrum Enterprise hides IPs. Shield Standard baseline.', 'DevOps'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 11. WHAT WE DON'T KNOW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. What We Do Not Know', level=1)
doc.add_paragraph('Significant unknowns that must be resolved by prototyping.')

for i, item in enumerate([
    'The actual single-server player ceiling for this game. No published data for a third-person MMORPG with hybrid combat and cosmetic customisation at 300+ in UE5.',
    'Replication Graph performance with CH\'s specific actor complexity. The no-stat cosmetic system means more visual state per character than typical.',
    'Actual ClientTravel transition time with full cosmetic state through Redis.',
    'MassAI client-side deterministic simulation viability. Floating-point determinism across diverse hardware is not guaranteed.',
    'If Replication Graph proves insufficient. No lower-risk path identified without reducing population, changing encounter design, or adopting licensed meshing.',
    'OmniMesh licensing terms. Pricing and source code access not publicly documented.',
    'Min-spec client rendering at 50+ nearby characters with CH\'s Byte-Punk art style.',
    'Custom metagame backend scalability at 30,000+ CCU. Must be designed for horizontal scaling from the start.',
], 1):
    add_bullet(item, bold_prefix=f'{i}.')

# ═══════════════════════════════════════════════════════════════
# 12. PHASED DELIVERY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Phased Delivery', level=1)

doc.add_heading('Phase 1: Foundation (Pre-production)', level=2)
doc.add_paragraph('Validate the architecture with hard data before committing production resources.')
add_bullet('200-player stress test prototype with Replication Graph and custom interest management node')
add_bullet('Shard transition prototype (ClientTravel + Redis round-trip with full cosmetic state)')
add_bullet('MassAI determinism test across 3+ hardware configurations')
add_bullet('OmniMesh licensing inquiry (only if prototype ceiling is below required target range)')
add_bullet('Custom metagame backend prototype (accounts, social graph, matchmaking) on Kubernetes')
doc.add_paragraph('Team: 2 to 3 networking engineers + 1 DevOps + 1 backend engineer, full-time.')

add_callout('PROTOTYPE GATE (all must pass)',
    'Median tick rate >= 15 Hz, p95 >= 12 Hz, p99 >= 10 Hz during 1-hour run. NetBroadcastTickTime < 40% of tick budget. Per-connection bandwidth < 50 KB/s steady state. No server-induced packet loss above controlled baseline. Reliable RPC queue depth stays bounded. Memory and GC stable (no leaks). 100-player combat cluster subjectively acceptable at 100ms RTT. Client rendering at 50+ nearby characters on min-spec: > 30 FPS. If any fail: reduce target and/or evaluate OmniMesh.',
    'CC3333')

doc.add_heading('Phase 2: Core Infrastructure (Early Production)', level=2)
add_bullet('Interest management with three tiers, hysteresis, combat elevation')
add_bullet('Hub sharding for Downtime with social-graph-aware routing')
add_bullet('State transfer pipeline (Redis + PostgreSQL + verification handshake)')
add_bullet('Instance server pipeline via Agones warm pool')
add_bullet('User Space housing server pipeline')
add_bullet('Corruption state replication (zone manager + central world-state service)')
add_bullet('Custom metagame backend core features')
add_bullet('Telemetry dashboard, EasyAntiCheat, Cloudflare Spectrum deployment')
doc.add_paragraph('Team: 2 to 3 networking + 1 to 2 DevOps + 1 to 2 backend.')

doc.add_heading('Phase 3: Scale and Polish (Mid to Late Production)', level=2)
add_bullet('300-player stress tests under production conditions')
add_bullet('Continuous stress testing in CI')
add_bullet('Server crash recovery validation')
add_bullet('Console compatibility (PSN/XBL session integration, TRC/XR validation)')
add_bullet('Min-spec client performance profiling')
add_bullet('All thresholds data-driven and tunable via config')

# ═══════════════════════════════════════════════════════════════
# 13. TEAM CAPABILITY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Team Capability Assessment', level=1)

add_table(
    ['Capability', 'Headcount', 'Details'],
    [
        ['UE5 dedicated server / networking', '2 to 3 engineers', 'Replication Graph internals, dormancy, NetDriver, server profiling. Custom C++ node development.'],
        ['Kubernetes / DevOps', '1 to 2 engineers', 'Agones deployment, CI/CD, monitoring, Redis/PostgreSQL ops, Cloudflare Spectrum.'],
        ['Backend development', '1 to 2 engineers', 'Custom metagame backend, world state service, shard routing logic.'],
        ['UE5 rendering (part-time)', '1 engineer', 'Impostor billboard pipeline, character LOD, animation LOD.'],
    ]
)

add_callout('STAFFING REQUIREMENT',
    'If the current tech team does not include 2 to 3 engineers with UE5 dedicated server experience, budget for hiring or contracting senior networking engineers. This is specialised work. Gameplay programmers cannot be retrained quickly enough.\n\nTotal infrastructure headcount: 5 to 7 during core build (2 to 3 networking, 1 to 2 DevOps, 1 to 2 backend). Reduces to 3 to 4 steady-state after launch stabilisation.',
    'CC3333')

# ═══════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_para('NBI Gaming / Confidential', bold=True, color=MUTED, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Prepared for Couch Heroes internal technical leadership', color=MUTED, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Figures cited from external sources represent their published claims, not independently verified measurements by NBI.', italic=True, color=MUTED, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Version 3.0 / June 2026', color=MUTED, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
out = r'd:\OneDrive\Claude_code\NBIAI_TEAM\Clients\Couch Heroes\technical\CH_Zone_Architecture_Design_Spec.docx'
doc.save(out)
print(f'Saved to {out}')
