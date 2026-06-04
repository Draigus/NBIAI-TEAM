"""Build CH Health and Safety Policy Word document with professional formatting."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC = Path(__file__).parent / "CH_Health_and_Safety_Policy.md"
DST = Path(__file__).parent / "CH_Health_and_Safety_Policy.docx"

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    if level == 1:
        h.font.size = Pt(20)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(15)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

# --- Read markdown ---
md = SRC.read_text(encoding='utf-8')
lines = md.split('\n')

def add_bold_text(paragraph, text):
    """Parse **bold** markers in text and add runs accordingly."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def format_table(table):
    """Apply consistent formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders = doc.element.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = doc.element.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:color'): '999999',
            qn('w:space'): '0',
        })
        borders.append(el)
    tblPr.append(borders)

    # Header row shading
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            shading = doc.element.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): '1B3A5C',
            })
            cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

i = 0
in_table = False
table_rows = []

while i < len(lines):
    line = lines[i]

    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue

    # Table detection
    if '|' in line and not in_table:
        # Check if this is a table (has | delimiters)
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]
        if cells:
            in_table = True
            table_rows = [cells]
            i += 1
            continue

    if in_table:
        if '|' in line:
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]
            # Skip separator rows (---|----|---)
            if cells and all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            if cells:
                table_rows.append(cells)
            i += 1
            continue
        else:
            # End of table — render it
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < max_cols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_bold_text(p, cell_text)
                format_table(table)
                doc.add_paragraph()
            in_table = False
            table_rows = []
            continue

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue

    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue

    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    # Numbered list items
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        add_bold_text(p, m.group(2))
        i += 1
        continue

    # Bullet list items
    if line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_bold_text(p, line[2:])
        i += 1
        continue

    # Empty lines
    if not line.strip():
        i += 1
        continue

    # Regular paragraphs (including bold field labels like **Name:**)
    p = doc.add_paragraph()
    text = line.strip()

    # Italic text
    if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
        run = p.add_run(text.strip('*'))
        run.italic = True
        run.font.size = Pt(10)
    else:
        add_bold_text(p, text)

    i += 1

# Flush any remaining table
if in_table and table_rows:
    max_cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=max_cols)
    for ri, row_data in enumerate(table_rows):
        for ci, cell_text in enumerate(row_data):
            if ci < max_cols:
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                add_bold_text(p, cell_text)
    format_table(table)

doc.save(str(DST))
print(f"Saved to {DST}")
