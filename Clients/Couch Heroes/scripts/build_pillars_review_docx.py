"""Build the Couch Heroes Design Pillars review and iteration DOCX from Markdown.

Deliberately plain: Word-native heading styles, no cover page, no colour blocks,
no branded furniture. The register matches Simon's own unified draft, because this
document is meant to be diffed against it and argued with, not presented.

Usage:
    python build_pillars_review_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, Inches

BASE = Path(__file__).resolve().parents[1] / "game_design"
SRC = BASE / "Couch_Heroes_Design_Pillars_Review_and_Iteration_2026-07-27.md"
OUT = BASE / "Couch_Heroes_Design_Pillars_Review_and_Iteration_2026-07-27.docx"

# Headings that should start on a fresh page.
PAGE_BREAK_BEFORE = re.compile(
    r"^(Part 1:|Part 2:|Pillar \d+:|The overriding design philosophy|Decisions for the session)"
)

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_runs(paragraph, text):
    """Render **bold** and *italic* inline markers into Word runs."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        else:
            paragraph.add_run(chunk)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in (
        ("Heading 1", 20, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 12, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)


def parse_table(lines, start):
    """Consume a pipe table starting at lines[start]. Returns (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    footer = section.footer.paragraphs[0]
    footer.text = "Couch Heroes design pillars: review and iteration. Draft for alignment, 27 July 2026."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    first_heading_done = False
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        cell = table.cell(r, c)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        para.paragraph_format.space_after = Pt(4)
                        add_runs(para, cell_text)
                        if r == 0:
                            for run in para.runs:
                                run.bold = True
                doc.add_paragraph()
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            content = stripped.lstrip("#").strip()
            # The document's own H1/H2 title block maps to Title/Subtitle.
            if not first_heading_done and level == 1:
                para = doc.add_paragraph(content, style="Title")
                first_heading_done = True
                i += 1
                continue
            style = f"Heading {min(level, 3)}"
            if PAGE_BREAK_BEFORE.match(content):
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            para = doc.add_paragraph(style=style)
            add_runs(para, content)
            i += 1
            continue

        if stripped.startswith("> "):
            para = doc.add_paragraph(style="Intense Quote")
            add_runs(para, stripped[2:])
            i += 1
            continue

        if stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, stripped[2:])
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, numbered.group(2))
            i += 1
            continue

        para = doc.add_paragraph()
        add_runs(para, stripped)
        i += 1

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
