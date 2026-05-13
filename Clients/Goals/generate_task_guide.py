"""
Generate Goals Studio Pricing Engagement Task Guide as a Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"d:\OneDrive\Claude_code\NBIAI_TEAM\Clients\Goals\Goals_Pricing_Engagement_Task_Guide.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Base styles ───────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading_style(style, size, bold=True, color=None, space_before=12, space_after=6):
    style.font.name = 'Calibri'
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)

set_heading_style(doc.styles['Heading 1'], 16, color=(0, 70, 127))
set_heading_style(doc.styles['Heading 2'], 14, color=(31, 73, 125))
set_heading_style(doc.styles['Heading 3'], 12, color=(68, 114, 196))

# ── Helper functions ──────────────────────────────────────────────────────────

def add_page_break(document):
    document.add_page_break()

def add_heading(document, text, level):
    return document.add_heading(text, level=level)

def add_para(document, text='', bold=False, italic=False, size=None, space_before=0, space_after=6):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_labelled(document, label, text):
    """Add a paragraph with a bold label followed by normal text."""
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(11)
    run_text = p.add_run(' ' + text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(11)
    return p

def add_numbered_list(document, items):
    """Add a numbered list from a list of strings (or sub-lists for indented bullets)."""
    for i, item in enumerate(items, 1):
        if isinstance(item, str):
            p = document.add_paragraph(style='List Number')
            p.add_run(item).font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
        elif isinstance(item, list):
            # Sub-bullets
            for sub in item:
                p = document.add_paragraph(style='List Bullet 2')
                p.add_run(sub).font.name = 'Calibri'
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

def add_bullet_list(document, items, indent=0):
    style = 'List Bullet' if indent == 0 else 'List Bullet 2'
    for item in items:
        p = document.add_paragraph(style=style)
        p.add_run(item).font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

def add_table(document, headers, rows):
    table = document.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    hdr.cells[0].paragraphs[0].paragraph_format.space_after = Pt(3)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            # Alternate row shading
            if r_idx % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    document.add_paragraph()  # spacing after table
    return table

def add_task_block(document, task_id, title, hours, what_text,
                   how_items, done_when_text):
    """Render a full task block."""
    h = document.add_heading(f'{task_id}: {title}', level=3)

    # Meta line
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r3 = p.add_run('Estimated Hours: ')
    r3.bold = True; r3.font.name = 'Calibri'; r3.font.size = Pt(11)
    r4 = p.add_run(str(hours))
    r4.font.name = 'Calibri'; r4.font.size = Pt(11)

    # What
    p2 = document.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(3)
    rw = p2.add_run('What: ')
    rw.bold = True; rw.font.name = 'Calibri'; rw.font.size = Pt(11)
    rw2 = p2.add_run(what_text)
    rw2.font.name = 'Calibri'; rw2.font.size = Pt(11)

    # How header
    ph = document.add_paragraph()
    ph.paragraph_format.space_before = Pt(4)
    ph.paragraph_format.space_after = Pt(2)
    rh = ph.add_run('How:')
    rh.bold = True; rh.font.name = 'Calibri'; rh.font.size = Pt(11)

    # How items: list of (text, [sub_bullets])
    for step_num, step in enumerate(how_items, 1):
        if isinstance(step, str):
            p_step = document.add_paragraph(style='List Number')
            p_step.add_run(step).font.name = 'Calibri'
            p_step.paragraph_format.space_before = Pt(1)
            p_step.paragraph_format.space_after = Pt(1)
        elif isinstance(step, tuple):
            # (main_text, [sub_bullets])
            p_step = document.add_paragraph(style='List Number')
            p_step.add_run(step[0]).font.name = 'Calibri'
            p_step.paragraph_format.space_before = Pt(1)
            p_step.paragraph_format.space_after = Pt(1)
            for sub in step[1]:
                p_sub = document.add_paragraph(style='List Bullet 2')
                p_sub.add_run(sub).font.name = 'Calibri'
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after = Pt(0)

    # Done When
    p3 = document.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    p3.paragraph_format.space_after = Pt(8)
    rd = p3.add_run('Done When: ')
    rd.bold = True; rd.font.name = 'Calibri'; rd.font.size = Pt(11)
    rd2 = p3.add_run(done_when_text)
    rd2.font.name = 'Calibri'; rd2.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(12)
r = p.add_run('Goals Studio Pricing Engagement')
r.font.name = 'Calibri'
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0, 70, 127)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run('Task Guide')
r2.font.name = 'Calibri'
r2.font.size = Pt(22)
r2.bold = True
r2.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()  # spacer

for label, value in [
    ('Prepared by:', 'NBI Analytics Ltd'),
    ('Date:', 'April 2026'),
    ('Deadline:', '27 April 2026'),
    ('Client:', 'Goals AB (Goals Studio)'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    rl = p.add_run(label + '  ')
    rl.font.name = 'Calibri'
    rl.font.size = Pt(12)
    rl.bold = True
    rv = p.add_run(value)
    rv.font.name = 'Calibri'
    rv.font.size = Pt(12)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Executive Summary', 1)

add_para(doc, (
    'NBI Analytics Ltd has been engaged by Goals AB (Goals Studio) to deliver an independent '
    'pricing review in advance of the game\'s May 14, 2026 commercial launch. The engagement '
    'covers Hard Currency (HC) pricing validation, competitive benchmarking against the Free-to-Play (F2P) '
    'sports market, regional pricing assessment across 40+ territories, and specific, '
    'actionable recommendations that Goals can implement before platform submission.'
), space_after=8)

add_para(doc, (
    'Goals\' internal team (Jonas, Julius) have already built a comprehensive 40-country '
    'pricing matrix with Purchasing Power Parity (PPP) columns, Foreign Exchange (FX) deviations, and competitive data across seven games. '
    'NBI\'s role is not to replicate that work. NBI\'s value is: (a) independent external '
    'validation providing confidence to submit to platforms, (b) insights that their internal '
    'position prevents them from seeing, and (c) specific recommendations with modelled impact '
    'ranges that require genre benchmarks unavailable in-house.'
), space_after=8)

add_heading(doc, 'Budget and Team Allocation', 2)

add_table(doc,
    ['Item', 'Detail'],
    [
        ['Total budget', '100,000 SEK (~$10,000 USD)'],
        ['Total estimated hours', '~116 hours'],
        ['Deadline', '27 April 2026 (client delivery)'],
        ['Review call', 'Week of 27 April 2026'],
    ]
)

add_heading(doc, 'Engagement Architecture', 2)

add_table(doc,
    ['Phase', 'Feature', 'Hours', '% of Budget'],
    [
        ['Foundation', 'F1: Competitive HC Pricing Benchmark', '38h', '33%'],
        ['Competitive benchmark', 'F2: Regional Pricing Strategy', '34h', '29%'],
        ['Regional analysis', 'F3: Risk Assessment, Scenarios and Recommendations', '16h', '14%'],
        ['Synthesis', 'F4: Value-Add (Framework and World Cup)', '8h', '7%'],
        ['Deliverable', 'F5: Deliverable Assembly and Client Review', '18h', '16%'],
        ['', 'Total', '~116h', '100%'],
    ]
)

add_para(doc, (
    'Feature 4 is conditional: it is only executed if Features 1-3 complete cleanly within '
    'budget. If any feature overruns, Feature 4 scope is reduced first. The core contracted '
    'deliverable (Features 1-3 plus Feature 5) totals approximately 106 hours.'
), space_after=8)

add_para(doc, (
    'Core delivery principle: Launch prices high; patch down if needed. Console platforms '
    'make post-launch price increases extremely difficult and commercially damaging. Every '
    'recommendation in this engagement reflects that constraint.'
), space_after=8)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 1
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Feature 1: Competitive HC Pricing Benchmark', 1)

add_para(doc, (
    'Objective: Show Goals where they sit relative to the sports/football F2P market. '
    'Extend their existing competitive analysis with insights their internal position does '
    'not give them.'
), space_after=4)

add_para(doc, (
    'What NBI adds that Goals cannot do internally: Independent cross-genre validation, '
    'pricing psychology assessment, conversion rate benchmarks derived from primary research '
    'of comparable live titles, and the confidence that an external expert agrees (or '
    'disagrees) with their conclusions.'
), space_after=4)

add_labelled(doc, 'Estimated feature budget:', '38 hours (33% of total engagement)')

doc.add_paragraph()

# ─── Story 1.0 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 1.0: Foundation -- Map the HC Economy', 2)
add_para(doc, (
    'Purpose: Before comparing prices, establish exactly what Goals\' HC buys and what the '
    'player value chain looks like. This is the prerequisite for everything else.'
), space_after=6)

# Task 1.0.1
add_task_block(
    doc,
    task_id='Task 1.0.1',
    title='Map HC to Value Chain',
    hours='3h',
    what_text=(
        'Document the complete path from real money to in-game value in Goals.'
    ),
    how_items=[
        ('From the Goals Studio Monetisation Design Doc and the Pricing Model data, establish:', [
            'Hard Currency (Coins): Purchased with real money only. Used for: kits, celebrations, and as an accelerator to soft currency (buy players, discard, get points).',
            'Soft Currency (Points): Earned through gameplay. Used for: player packs, upgrades, legend conversions, marketplace.',
            'Conversion: 25,000 points = 500 coins = $6.00 USD (from Pricing Model, Sheet: Conversion rates).',
            'Therefore: 1 HC coin = 50 points = $0.012 USD at base rate.',
        ]),
        ('Map what each HC tier actually buys:', [
            'Source for item pricing: Goals\' Content Plan (Pricing Model spreadsheet) shows Internal Kits at 16,625 points ($3.99), Partnership Kits at 29,125 points ($6.99), Celebrations at 29,167 points ($7.00).',
            'Convert to HC: Internal Kit = 16,625 / 50 = ~333 coins. Partnership Kit = 29,125 / 50 = ~583 coins. Celebration = 29,167 / 50 = ~583 coins.',
            '380 HC ($5.99): buys 1 Internal Kit (333 coins) with 47 coins left over. Nothing else affordable.',
            '650 HC ($9.99): buys 1 Internal Kit (333 coins) plus remainder of 317 coins (not enough for a second kit or celebration).',
            'GAP: Mech and Pulse kit HC pricing is NOT documented in any client material. The beta sold these for soft currency only. Their HC pricing at launch needs to be confirmed with Julius.',
        ]),
        ('Map item categories and their HC costs (from Pricing Model Content Plan):', [
            'Internal Kits: ~333 HC (source: 16,625 pts / 50).',
            'Partnership Kits: ~583 HC (source: 29,125 pts / 50).',
            'Celebrations: ~583 HC (source: 29,167 pts / 50).',
            'Player packs (via HC to points): varies by tier -- see Standard Packs in Pricing Model.',
            'UNRESOLVED: Specific HC prices for Mech, Pulse, National kit categories. These tiers appear in beta community data but without confirmed launch pricing.',
        ]),
        'Key finding: Document the "minimum meaningful purchase" -- what is the cheapest satisfying thing a new player can buy with HC? Currently the entry tier (380 HC / $5.99) buys exactly 1 Internal Kit with minimal leftover. This is a tight fit -- any price increase at the entry tier could push the cheapest item out of reach.',
    ],
    done_when_text=(
        'Complete HC value chain documented. Every store item category mapped to HC cost with '
        'documented source. Gaps flagged for Julius. "Minimum meaningful purchase" identified '
        'and costed.'
    )
)

# Task 1.0.2
add_task_block(
    doc,
    task_id='Task 1.0.2',
    title='Establish Payer Persona Benchmarks',
    hours='5h',
    what_text=(
        'Using Goals\' 3-persona model (from their Pricing Model) and industry benchmarks, '
        'establish what healthy conversion and Average Revenue Per Paying User (ARPPU) look like for this game.'
    ),
    how_items=[
        ('Goals\' personas (from their Pricing Model, Sheet: Overview):', [
            'Persona 1 (Core FIFA Ultimate Team (FUT) player -- every day): 20% of players, 30% pay, ARPPU $46.',
            'Persona 2 (Dedicated -- 2-3x/week): 35% of players, 10% pay, ARPPU $15.81.',
            'Persona 3 (Casual -- few times/month): 45% of players, 2% pay, ARPPU $7.24.',
            'Overall payer rate: 10.4% | Overall Average Revenue Per User (ARPU): $3.38 | Overall ARPPU: $32.48.',
        ]),
        ('Benchmark these against verifiable sources:', [
            'Research method: Pull published conversion and ARPPU data from eFootball public earnings (Konami quarterly reports disclose mobile/console revenue and Daily Active Users (DAU) -- derive ARPPU), EA FC 25/26 public filings (EA quarterly earnings disclose Ultimate Team revenue and player counts), Sensor Tower / data.ai reports on F2P sports category (if accessible), Game Developers Conference (GDC) talks with published benchmarks (search for F2P monetisation presentations with cited data).',
            'What we are comparing: Goals\' 10.4% payer rate against actual published data from comparable titles -- NOT against vague "industry norms".',
            'NOTE: Goals\' own Pricing Model (Sheet: Context) includes a "Gemini deep research check" with alternate figures: Persona 1 paying 30% with ARPPU $25-$40, Persona 2 paying 15% with ARPPU $10-$15, Persona 3 paying 3% with ARPPU $5-$10. These are lower than Goals\' model assumes. Source of Gemini figures needs to be verified.',
        ]),
        'Assess: Is Goals\' 10.4% overall payer rate realistic for launch? Compare against the published data gathered above. If published F2P sports payer rates are materially lower, Goals\' revenue model overstates early revenue and the pricing strategy needs to account for a slower conversion ramp.',
        'Critical question: Is 10.4% a month-1 assumption or a steady-state target? The plan and revenue model do not specify the timeframe. If it is steady-state applied to month 1, the model overstates early revenue. This needs clarification from Jonas/Julius.',
        'Seasonless economy implication: Goals has no seasonal resets, no battle pass, no Fear of Missing Out (FOMO)-driven wipes. This means their HC pack ladder must do ALL the conversion work that seasonal games distribute across time-limited events. Compare directly: EA FC resets squads annually (forcing re-spend); eFootball is non-seasonal (closer model to Goals). In non-seasonal economies, the HC ladder needs to be more compelling at each tier because there is no artificial urgency driving purchases. This should factor into whether their discount curve is aggressive enough to drive ongoing conversion without seasonal pressure.',
    ],
    done_when_text=(
        'Goals\' payer assumptions benchmarked against published data from comparable titles '
        '(with cited sources). Clear assessment of whether their conversion/ARPPU targets are '
        'realistic at launch. If unrealistic, document the implication for pricing strategy.'
    )
)

# ─── Story 1.1 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 1.1: Competitive Pricing Position', 2)
add_para(doc, (
    'Purpose: Show Goals exactly where their HC pricing sits relative to the market -- not '
    'by re-doing their competitive data, but by synthesising it into a clear positioning '
    'statement they cannot write themselves.'
), space_after=6)

# Task 1.1.1
add_task_block(
    doc,
    task_id='Task 1.1.1',
    title='Build Normalised Price Position Map',
    hours='8h',
    what_text=(
        'Using Goals\' existing competitive data (which they have already gathered for 7 '
        'games), create a normalised view that shows positioning at each tier.'
    ),
    how_items=[
        'Goals already has the raw data in their Pricing Model (Sheet "Softhard currency comparison"). NBI\'s job is NOT to re-gather this data. NBI\'s job is to interpret it.',
        ('Create a "position map" for each tier showing Goals\' effective USD/HC against the competitive range:', [
            'At entry ($5-6): Goals = $0.01576/HC. Where does this sit vs range?',
            'At mid ($20): Goals = $0.01454/HC. Position?',
            'At top ($100): Goals = $0.01290/HC. Position?',
        ]),
        'For each tier, calculate market: minimum, median, maximum, and Goals\' position as percentile.',
        'Visual: Goals should see "You are at the Xth percentile of the market at each tier" -- this is the answer to "are we too low?"',
        'Spot-check 2-3 competitor prices to confirm Goals\' data is still current (quick store verification, not full re-research). If EA FC or Fortnite changed prices since Goals\' data was gathered, note the update.',
    ],
    done_when_text=(
        'Position map complete showing Goals\' percentile at each tier. Clear one-line answer: '
        '"Goals is at the [X]th percentile of market pricing -- [below/at/above] median."'
    )
)

# Task 1.1.2
add_task_block(
    doc,
    task_id='Task 1.1.2',
    title='Discount Curve Comparison',
    hours='4h',
    what_text=(
        'Compare Goals\' bonus/discount scaling against competitors. Assess whether the value '
        'incentive to buy larger packs is appropriate.'
    ),
    how_items=[
        ('Goals\' discount curve (already calculated from their data):', [
            'Tier 1 ($5.99): 0% discount (base).',
            'Tier 2 ($9.99): ~2.5% discount.',
            'Tier 3 ($19.99): ~7.7%.',
            'Tier 4 ($29.99): ~9.4%.',
            'Tier 5 ($49.99): ~15.4%.',
            'Tier 6 ($99.99): ~18.1%.',
        ]),
        ('Compare max discount to competitors:', [
            'EA FC: ~39% at top tier (very aggressive).',
            'Fortnite: ~24%.',
            'Apex: ~13%.',
            'Goals: ~18%.',
        ]),
        ('Assessment questions:', [
            'Is 18% max discount enough to incentivise whale spending at the top tier?',
            'Is the jump from tier to tier consistent enough to feel "fair"?',
            'Is there a clear "best value" tier that most players should gravitate toward?',
            'Compared to the market: Goals is between Apex (conservative) and Fortnite (moderate). This is a defensible position for a new launch -- do not over-discount before you have data on player willingness.',
        ]),
        'Key principle: New launches should be MORE conservative on discounts (closer to Goals\' current curve) because you can always increase discounts later but can never reduce them without backlash.',
    ],
    done_when_text=(
        'Curve comparison documented. Assessment of whether Goals\' discount progression is '
        'appropriate for a new launch. Specific recommendation if any tier\'s discount '
        'should change.'
    )
)

# Task 1.1.3
add_task_block(
    doc,
    task_id='Task 1.1.3',
    title='Price Tier (SKU) Structure and Entry Point Audit',
    hours='5h',
    what_text=(
        'Assess the structural health of Goals\' 6-tier HC ladder. Special focus on the '
        'entry point and first-purchase experience.'
    ),
    how_items=[
        ('Goals\' current structure:', [
            '$5.99 / $9.99 / $19.99 / $29.99 / $49.99 / $99.99 (6 tiers).',
        ]),
        ('Compare tier count to competitors:', [
            'EA FC: 8 tiers (starts at $0.99).',
            'Fortnite: 4 tiers (starts at $7.99).',
            'Apex: 6 tiers (starts at $4.99).',
        ]),
        ('Entry point analysis:', [
            'Goals starts at $5.99. EA FC starts at $0.99. Apex at $4.99.',
            'Question: Does $5.99 create first-purchase friction?',
            'From Task 1.0.1: What can $5.99 (380 HC) buy? If the answer is "1 national kit with leftover coins that cannot buy anything else," that is a satisfying entry purchase. If 380 HC cannot buy any complete item, the entry tier has a conversion problem.',
            'Counter-argument: Lower entry tiers ($0.99-$2.99) have high platform fee overhead (30% of $0.99 = $0.30 net revenue). For a small studio, the transaction cost may not justify it.',
        ]),
        ('Structural gaps:', [
            'Gap between $9.99 and $19.99 (100% price jump) -- is there a missing $14.99 tier?',
            'Gap between $49.99 and $99.99 (100% price jump) -- should there be a $69.99 or $79.99?',
            'These gaps are standard in the industry (Fortnite, Apex have similar jumps). Not necessarily a problem.',
        ]),
        ('Whale ceiling:', [
            '$99.99 is the top tier. Whales who want to spend more must purchase multiple times.',
            'Is this intentional? (Yes -- it avoids a single $250 "whale pack" that looks predatory.)',
            'What is the maximum a whale could spend per day/week? (No purchase limits documented in the game design docs.)',
            'Recommendation: No purchase limits + $99.99 ceiling = whale-friendly without optically aggressive pricing. This is correct for their anti-pay-to-win (P2W) positioning.',
        ]),
    ],
    done_when_text=(
        'Structure audit complete. Entry point assessment with recommendation. Whale ceiling '
        'analysis documented. Any structural gaps identified with recommendation (add tier / '
        'leave as-is / remove tier).'
    )
)

# Task 1.1.4
add_task_block(
    doc,
    task_id='Task 1.1.4',
    title='Platform Fee Impact and Net Revenue Table',
    hours='2h',
    what_text=(
        'Calculate actual net revenue per SKU per platform. Identify any margin issues.'
    ),
    how_items=[
        ('Build the following net revenue table:', []),
    ],
    done_when_text=(
        'Net revenue table complete. Any margin concerns flagged. Epic revenue advantage '
        'quantified.'
    )
)

# Add the net revenue table inline after the how step
add_table(doc,
    ['SKU', 'Gross', 'Steam Net (30%)', 'Epic Net (12%)', 'PS/Xbox Net (30%)'],
    [
        ['380 HC', '$5.99', '$4.19', '$5.27', '$4.19'],
        ['650 HC', '$9.99', '$6.99', '$8.79', '$6.99'],
        ['1,375 HC', '$19.99', '$14.00', '$17.59', '$14.00'],
        ['2,100 HC', '$29.99', '$21.00', '$26.39', '$21.00'],
        ['3,750 HC', '$49.99', '$35.00', '$43.99', '$35.00'],
        ['7,750 HC', '$99.99', '$70.00', '$87.99', '$70.00'],
    ]
)

add_bullet_list(doc, [
    'Assessment: All tiers are well above minimum viable margin. The $5.99 entry netting $4.19 is healthy (unlike a hypothetical $0.99 tier netting $0.69).',
    'Epic advantage: Goals nets ~26% more per transaction on Epic vs Steam/PS/Xbox at the same price ($5.27 vs $4.19 on a $5.99 item). This could inform promotional strategy (drive purchases through Epic where possible).',
    'Note for regional analysis: These are USD nets. Regional pricing on consoles uses "wholesale" rates that may differ from simply applying 30% to local prices. Goals\' matrix already has PS wholesale prices -- verify alignment.',
])
doc.add_paragraph()

# ─── Story 1.2 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 1.2: Conversion Rate and Elasticity Context', 2)
add_para(doc, (
    'Purpose: Provide the conversion/revenue modelling context that turns a static price '
    'comparison into actionable strategy.'
), space_after=6)

# Task 1.2.1
add_task_block(
    doc,
    task_id='Task 1.2.1',
    title='Price Elasticity Scenario Modelling',
    hours='8h',
    what_text=(
        'Model what happens to projected revenue if Goals adjusts prices +/- 10-20% '
        'at key tiers.'
    ),
    how_items=[
        ('Use Goals\' own revenue model (Pricing Model, Sheet "Overview"):', [
            '805,250 Monthly Active Users (MAU), 10.4% payer rate, $32.48 ARPPU = $2.72M total revenue.',
        ]),
        ('Model four scenarios:', [
            'Scenario A: Raise all prices 10%. Assume conversion drops 5-15% (elasticity assumption for F2P games: moderate). Net revenue change: +10% price x (100% minus X% conversion drop) = range.',
            'Scenario B: Raise all prices 20%. Assume conversion drops 10-25%. Net revenue change: calculate range.',
            'Scenario C: Raise entry tier only (from $5.99 to $6.99). Assume entry conversion drops 5-10% but mid/high tiers unaffected. Revenue impact: primarily affects first-time buyers.',
            'Scenario D: Keep prices, steepen discount curve. Same base price, but top tier gets 25% discount instead of 18%. Attracts more whale spending at top, neutral on entry conversion.',
        ]),
        'Present as impact ranges (conservative/optimistic), not point estimates.',
        'Key insight to surface: "If your conversion rate target (10.4%) is optimistic for month 1, lowering prices will not help -- you need time and content to build conversion. Therefore, launching at a slightly higher price point and reducing later is the correct strategy."',
        'Core pricing principle (from Glen, stated in kickoff call): Launch high, patch down. Never the reverse. Console platforms make price increases nearly impossible post-launch.',
    ],
    done_when_text=(
        '4 scenarios modelled with revenue impact ranges. Clear recommendation on which '
        'scenario maximises launch revenue while preserving future flexibility.'
    )
)

# ─── Story 1.3 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 1.3: Promotional and First-Purchase Pricing', 2)
add_para(doc, (
    'Purpose: Address the gap in a static HC pack strategy. First-purchase incentives and '
    'promotional pricing drive early conversion that the standard ladder alone cannot achieve.'
), space_after=6)

# Task 1.3.1
add_task_block(
    doc,
    task_id='Task 1.3.1',
    title='First-Purchase and Starter Pack Recommendations',
    hours='5h',
    what_text=(
        'Provide recommendations on first-time buyer incentives that drive conversion '
        'without discounting the core HC ladder.'
    ),
    how_items=[
        ('Industry standard: Most F2P games offer a one-time "Starter Pack" or "First Purchase Double" that dramatically increases first-purchase conversion:', [
            'EA FC: "Welcome Pack" and various one-time offers.',
            'Fortnite: "Starter Pack" ($3.99, cosmetics + 600 V-Bucks).',
            'Apex: "Champion Edition" and "Starter Pack".',
        ]),
        'Goals\' current approach: From Beta LiveOps data, they had a "Starter Pack" at 15,000 points (soft currency, not HC). No HC-specific first-purchase offer documented.',
        ('Recommendations to include:', [
            'First Purchase Bonus: Double HC on first-ever purchase at any tier. Research needed: verify what EA FC, Fortnite, Apex, eFootball currently offer as first-purchase incentives and what published data exists on conversion uplift. Do NOT cite specific conversion improvement percentages without a source.',
            'Limited Starter Pack: A one-time $3.99 offer (below the normal $5.99 entry) containing 500 HC + 1 Rare+ player + 1 exclusive kit. Creates urgency, breaks the first-purchase barrier, and introduces players to the HC economy.',
            'World Cup Launch Bundle (if timing allows): Tied to June event, limited availability, higher value per dollar than standard packs. Creates urgency without permanently discounting.',
        ]),
        'Framing: These are NOT changes to the core HC ladder (which is what the Statement of Work (SOW) covers). These are complementary offers that improve conversion on the core ladder. Position as value-add observation.',
    ],
    done_when_text=(
        '2-3 specific first-purchase recommendations documented with rationale and expected '
        'conversion impact. Clearly framed as complementary to the core pricing, not scope creep.'
    )
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 2
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Feature 2: Regional Pricing Strategy', 1)

add_para(doc, (
    'Objective: Validate Goals\' 40+ country pricing matrix. NBI adds competitive regional '
    'context, PPP-adjusted analysis, and platform compliance verification they cannot '
    'get from their own spreadsheet.'
), space_after=4)

add_para(doc, (
    'What NBI adds: Goals built their matrix using exchange rates and some PPP data. They '
    'cannot benchmark their regional pricing against what EA FC actually charges in Turkey '
    'or Brazil -- that requires platform store access or competitor intelligence. NBI '
    'provides this.'
), space_after=4)

add_labelled(doc, 'Estimated feature budget:', '34 hours (29% of total engagement)')
doc.add_paragraph()

# ─── Story 2.1 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 2.1: Regional Deviation and Risk Identification', 2)
add_para(doc, (
    'Purpose: Flag countries where Goals\' pricing creates material risk (too high = player '
    'backlash, too low = revenue leakage, misaligned = platform rejection).'
), space_after=6)

# Task 2.1.1
add_task_block(
    doc,
    task_id='Task 2.1.1',
    title='Categorise All Countries by Risk Level',
    hours='8h',
    what_text=(
        'Using Goals\' existing deviation data (already calculated in their matrix), '
        'categorise every country into green/amber/red.'
    ),
    how_items=[
        'Goals\' Pricing Matrix already calculates "Difference to currency exchange" for each country. Use this directly -- do not recalculate.',
        ('Categorise using deviation thresholds (defined for this engagement based on Goals\' pricing matrix data patterns and the "Recommendation" column Goals themselves already populated):', [
            'Green (proceed as-is): Deviation <15% from FX AND Goals\' own matrix marks "YES".',
            'Amber (review recommended): Deviation 15-25% OR Goals\' matrix marks "!" or "?" OR misaligned with competitor pricing in that market.',
            'Red (change required): Deviation >25% OR Goals\' matrix explicitly recommends a price move OR large player market (top 10 beta communicators) with >15% over-pricing.',
        ]),
        ('From Goals\' data, expected categorisation:', [
            'Red candidates: Poland (+25%), Switzerland (+26%), Ukraine (-26%), possibly Turkey, Brazil.',
            'Amber candidates: South Africa, Czech Republic, Norway, Denmark.',
            'Green: Most of Western Europe (EUR countries at +18% are borderline amber).',
        ]),
        ('For each Red country, assess:', [
            'How large is the player base? (Cross-reference with Beta community data: UK 1,274, France 1,026, Turkey 667, Brazil 495, Poland 392.)',
            'What is the risk severity? (Large player base + over-priced = high severity.)',
        ]),
        'Special attention to the EUR zone: Goals prices EUR at the same level across all EUR countries (EUR5.99 base). But PPP within the EUR zone varies significantly (Germany vs Portugal vs Greece). This is standard practice (platform limitations), but flag if any EUR country has extreme PPP deviation.',
    ],
    done_when_text=(
        'Every country categorised (green/amber/red). Red countries have specific risk '
        'descriptions. Output is a ranked priority list for deeper analysis.'
    )
)

# Task 2.1.2
add_task_block(
    doc,
    task_id='Task 2.1.2',
    title='Competitor Regional Pricing for Red/Amber Markets',
    hours='15h',
    what_text=(
        'For the 5-10 countries flagged as red/amber, determine what EA FC and Fortnite '
        'charge. This is the intelligence Goals cannot access internally.'
    ),
    how_items=[
        ('For each red/amber country, research:', [
            'EA FC equivalent base tier pricing on PS Store and Steam in that region.',
            'Fortnite V-Bucks pricing in that region (Epic Games Store).',
        ]),
        ('Sources:', [
            'PlayStation Store regional pricing: Accessible via web store with region selection.',
            'Steam: SteamDB provides historical and current pricing by region.',
            'Epic: Regional pricing visible through store.',
        ]),
        ('For each market, build a comparison table:', [
            'Country | Goals proposed | EA FC actual | Fortnite actual | Market median | Goals position (above/below/at).',
        ]),
        ('Focus on the highest-risk markets first:', [
            'Brazil: Large player base, low PPP, critical for World Cup engagement.',
            'Turkey: 5th largest beta community, extreme PPP differential.',
            'Poland: Large beta community, Goals is +25% above FX.',
            'Russia/CIS: If launching there -- extreme currency volatility.',
        ]),
        'Note: If direct store access is blocked by region-locking, use SteamDB or documented third-party sources. Cite all sources.',
    ],
    done_when_text=(
        'Competitor pricing gathered for all red/amber countries. Goals\' position vs EA FC '
        'documented per market. Specific countries where Goals is significantly above or '
        'below EA FC identified. (Note: This is the single largest research task in the '
        'engagement.)'
    )
)

# ─── Story 2.2 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 2.2: Platform Compliance and Submission Guidance', 2)
add_para(doc, (
    'Purpose: Ensure Goals\' pricing will pass PS/Xbox review without rejection or delay.'
), space_after=6)

# Task 2.2.1
add_task_block(
    doc,
    task_id='Task 2.2.1',
    title='Verify Platform Tier Compliance',
    hours='3h',
    what_text=(
        'Confirm that Goals\' proposed prices for PS and Xbox align with valid platform '
        'tier selections.'
    ),
    how_items=[
        ('Goals\' matrix already has PS wholesale and Xbox pricing columns. Verify these are valid tier selections by cross-referencing with:', [
            'Goals\' own "Copy of PS4_PS5 Digital" sheet in the Pricing Matrix.',
            'Goals\' "Xbox wholesale pricing" sheet.',
        ]),
        'For any price that does not align to a standard platform tier, flag it -- Sony/Xbox will reject non-standard pricing.',
        'Key question for Julius: Have they already submitted a test pricing configuration? If yes, did it pass? If no, this verification is critical path.',
        ('Document the submission checklist:', [
            'All prices must be in valid platform tiers.',
            'Regional pricing must cover all territories where the game is available.',
            'Any currency not supported by the platform needs a USD fallback.',
            'Review period: ~9 days (from Julius\'s kickoff call statement).',
        ]),
        'Timeline implication: If NBI delivers recommendations by Sunday April 27, Goals has until approximately May 5 to submit. That gives them 1+ week to implement changes and submit.',
    ],
    done_when_text=(
        'Platform compliance verified or issues flagged. Submission timeline documented. '
        'Any blockers escalated to Julius immediately.'
    )
)

# ─── Story 2.3 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 2.3: Country-Specific Recommendations', 2)
add_para(doc, (
    'Purpose: Provide specific price adjustment recommendations for each red/amber country.'
), space_after=6)

# Task 2.3.1
add_task_block(
    doc,
    task_id='Task 2.3.1',
    title='Write Country Recommendation Cards',
    hours='8h',
    what_text=(
        'For each red/amber country, write a "recommendation card" with the specific '
        'pricing change needed.'
    ),
    how_items=[
        ('For each country requiring action, document the following fields:', [
            'Country: [Name].',
            'Current proposed price: [X local currency] (= $Y USD equivalent).',
            'Issue: [Over-priced by Z% vs FX / Misaligned with EA FC by W% / PPP unaffordable].',
            'Recommended price: [A local currency] (= $B USD equivalent).',
            'Rationale: [EA FC charges C in this market / PPP suggests D / Platform tier constraint means E].',
            'Revenue impact: [If player base is ~N, this change affects ~$M of annual revenue].',
            'Risk if unchanged: [Player perception / Platform rejection / Competitive disadvantage].',
        ]),
        'Prioritise cards by: player base size multiplied by deviation severity.',
        ('Group by action type:', [
            '"Reduce price" (over-priced markets).',
            '"Can increase" (under-priced markets where revenue is being left on the table).',
            '"Platform tier adjustment" (price needs rounding to valid tier).',
        ]),
        'Expected output: 8-12 country cards covering all red and amber markets.',
    ],
    done_when_text=(
        'All red/amber countries have specific, actionable recommendation cards. Each card '
        'is self-contained -- Julius can implement it without reading the full report.'
    )
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 3
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Feature 3: Risk Assessment, Scenarios and Recommendations', 1)

add_para(doc, (
    'Objective: Synthesise all analysis into clear, specific, actionable guidance with '
    'modelled impact.'
), space_after=4)

add_labelled(doc, 'Estimated feature budget:', '16 hours (14% of total engagement)')
doc.add_paragraph()

# ─── Story 3.1 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 3.1: "Are We Too Low?" -- The Central Answer', 2)
add_para(doc, (
    'Purpose: Directly address Jonas\'s core concern with evidence.'
), space_after=6)

# Task 3.1.1
add_task_block(
    doc,
    task_id='Task 3.1.1',
    title='Build the Definitive Position Statement',
    hours='4h',
    what_text=(
        'Using all Feature 1 analysis, write the one paragraph that answers Jonas\'s '
        'question: "Are we pricing too low?"'
    ),
    how_items=[
        'From Task 1.1.1: Goals\' percentile position at each tier (below/at/above median).',
        'From Task 1.0.2: Whether Goals\' conversion/ARPPU targets are realistic.',
        'From Task 1.2.1: Scenario modelling showing impact of price changes.',
        ('Synthesise into a clear statement:', [
            '"Goals\' base HC pricing is at the [X]th percentile of the competitive set. This means [interpretation]. Given your 80/20 skill-to-purchase ratio and anti-P2W positioning, this is [appropriate/low/high] because [reason]."',
            '"Your payer conversion target of 10.4% is [realistic/aggressive] for month 1. At realistic month-1 conversion of [Y%], your current pricing yields [Z revenue]. At 10% higher pricing with [slightly reduced] conversion, you would yield [W revenue]."',
        ]),
        'Conclude with a specific recommendation: e.g., "maintain current base pricing but steepen the discount curve at top tier" OR "increase base tier from $5.99 to $6.49 across all platforms".',
        'Core principle (from kickoff call): "You can always reduce prices post-launch. You cannot raise them." This is a platform constraint, not just strategy -- Sony/Xbox re-certification for price increases is slow and generates press coverage.',
    ],
    done_when_text=(
        'One clear paragraph answering "are we too low?" with specific evidence and a '
        'specific recommendation. No hedging -- take a position.'
    )
)

# ─── Story 3.2 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 3.2: Risk Register', 2)
add_para(doc, (
    'Purpose: Document all identified pricing risks before launch, ranked by severity. '
    '(Full risk register also appears as Appendix A.)'
), space_after=6)

# Task 3.2.1
add_task_block(
    doc,
    task_id='Task 3.2.1',
    title='Compile and Rank Risks',
    hours='5h',
    what_text=(
        'Full risk register from all analysis. Each risk documented with: Description, '
        'Severity, Likelihood, Markets Affected, and Mitigation.'
    ),
    how_items=[
        'Collect all risks identified across Features 1 and 2 analysis.',
        'Rank by severity x likelihood composite score.',
        ('The 10 expected risks (see Appendix A for full detail):', [
            'Risk 1: P2W perception at high price points.',
            'Risk 2: Post-launch price increase impossibility.',
            'Risk 3: Over-pricing in price-sensitive markets (Turkey, Brazil, Poland).',
            'Risk 4: Platform pricing rejection.',
            'Risk 5: Discount curve too shallow for launch.',
            'Risk 6: Missing entry-tier conversion.',
            'Risk 7: Originals/Celebrity pack pricing gap.',
            'Risk 8: Netherlands regulatory risk for player packs.',
            'Risk 9: South Korea drop rate disclosure requirement.',
            'Risk 10: Non-seasonal demand sustainability.',
        ]),
        'Top 3 risks must have specific, actionable mitigation steps.',
    ],
    done_when_text=(
        'Minimum 10 risks documented with full detail. Ranked by severity x likelihood. '
        'Top 3 have specific, actionable mitigation steps.'
    )
)

# ─── Story 3.3 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 3.3: Top 5 Recommendations', 2)
add_para(doc, (
    'Purpose: The five highest-impact actions Goals should take before platform submission.'
), space_after=6)

# Task 3.3.1
add_task_block(
    doc,
    task_id='Task 3.3.1',
    title='Write Top 5 Recommendations',
    hours='7h',
    what_text=(
        'Five specific, numbered, actionable recommendations with full supporting detail.'
    ),
    how_items=[
        ('Each recommendation must include:', [
            'What to change: Exact specification.',
            'Why: Evidence from analysis.',
            'Revenue impact: Conservative-to-optimistic range (from scenario modelling).',
            'Implementation: What Goals does (which cells in their matrix, which platform to resubmit).',
            'Risk if skipped: What happens.',
        ]),
        ('The 5 recommendations will be determined by analysis, but expected candidates:', [
            'Recommendation 1 -- Pricing position adjustment: e.g., "Increase/maintain base tier at $X" based on competitive position finding.',
            'Recommendation 2 -- Regional correction: e.g., "Reduce Poland/Turkey/Brazil pricing by X%" based on country analysis.',
            'Recommendation 3 -- Discount curve adjustment: e.g., "Increase top-tier discount from 18% to 22%" if whale incentive is weak.',
            'Recommendation 4 -- First-purchase incentive: e.g., "Implement first-purchase double bonus" to drive conversion.',
            'Recommendation 5 -- Launch framing: e.g., "Label current pricing as Launch Pricing to preserve future flexibility".',
        ]),
        'Each must pass the test: "Would a studio economy designer respect this and act on it immediately?"',
    ],
    done_when_text=(
        '5 recommendations written in full format. No vague advice. Every number is '
        'specific. Every impact range is stated.'
    )
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 4
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Feature 4: Value-Add (Conditional -- Time Permitting)', 1)

add_para(doc, (
    'Objective: Deliver supplementary tools that demonstrate NBI\'s strategic depth and '
    'position Phase 2. ONLY executed if Features 1-3 complete within budget.'
), space_after=4)

add_labelled(doc, 'Estimated feature budget:', '8 hours (7% of total engagement) -- conditional')
doc.add_paragraph()

# ─── Story 4.1 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 4.1: Reusable Regional Pricing Methodology', 2)

# Task 4.1.1
add_task_block(
    doc,
    task_id='Task 4.1.1',
    title='Write the "How to Price a New SKU" Guide',
    hours='3h',
    what_text=(
        'One-page methodology Goals can use for any future SKU without NBI.'
    ),
    how_items=[
        ('Document the steps NBI used in this engagement, simplified for internal use:', [
            'Step 1: Set USD base price (anchor against EA FC equivalent).',
            'Step 2: Convert to local currencies using current FX.',
            'Step 3: Check PPP factor (provide lookup table).',
            'Step 4: Compare against EA FC in that market (+/-15% is acceptable).',
            'Step 5: Round to nearest platform tier.',
            'Step 6: Verify discount curve is maintained.',
        ]),
        'Include one worked example: "Pricing a new Celebration at $7.00 USD for the Brazil market".',
        'Include the PPP factor table for all launch markets.',
        'Position as: "This is the manual method. Phase 2 builds a model that does this automatically using your live data."',
    ],
    done_when_text=(
        'One-page document that Julius can follow. Worked example included.'
    )
)

# ─── Story 4.2 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 4.2: World Cup Pricing Opportunity Notes', 2)

# Task 4.2.1
add_task_block(
    doc,
    task_id='Task 4.2.1',
    title='Write World Cup Quick-Hit Observations',
    hours='2h',
    what_text=(
        '3-5 tactical pricing observations for the June World Cup window.'
    ),
    how_items=[
        ('From the Release LiveOps data, Goals already plans:', [
            'National kits by country.',
            'World Cup themed events.',
            'Peak monetisation in July.',
        ]),
        ('NBI adds:', [
            'National bundle pricing guidance: Price national kits at a premium during group stage (emotional spending) and discount during knockout (losing teams = fire sale).',
            'Time-limited HC pack: Consider a "World Cup Coin Pack" at slightly better value than standard (e.g., 10% bonus) that is only available during the tournament -- drives urgency.',
            'ARPPU peak expectations: Research what published data exists on revenue spikes during major sporting events in comparable F2P titles (EA FC during World Cup 2022, eFootball during tournaments). Do NOT cite multiplier ranges without a verifiable source.',
        ]),
        'Keep this SHORT. Half a page. It is a teaser for the live service consulting engagement, not a full plan.',
    ],
    done_when_text=(
        '3-5 bullet points with specific, actionable World Cup pricing tactics. Clearly '
        'positioned as "what we would build out fully in Phase 2."'
    )
)

# ─── Story 4.3 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 4.3: Phase 2 Positioning', 2)

# Task 4.3.1
add_task_block(
    doc,
    task_id='Task 4.3.1',
    title='Write "What\'s Next" Section',
    hours='3h',
    what_text=(
        'Brief section positioning the larger engagement. Frame as recommendations, '
        'not a sales pitch.'
    ),
    how_items=[
        ('Natural next steps:', [
            'Post-launch telemetry review (2-3 weeks after May 14): Real data replaces estimates.',
            'A/B test design: Optimise conversion at each tier using live cohorts.',
            'Store optimisation: Offer sequencing, time-limited promotion design, personalised pricing.',
            'World Cup monetisation strategy: Full tactical plan for the June peak.',
            'Ongoing live service consulting: Content cadence, engagement frameworks, team coaching.',
        ]),
        'Close with: "We recommend a check-in call 2-3 weeks post-launch to review initial transaction data and scope the next phase."',
        'Do NOT include NBI\'s internal pricing for Phase 2 in the client document.',
    ],
    done_when_text=(
        'Half-page "What\'s Next" section. Helpful, not salesy. Positions naturally.'
    )
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 5
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Feature 5: Deliverable Assembly and Client Review', 1)

add_para(doc, (
    'Objective: Package everything into the contracted deliverable.'
), space_after=4)

add_labelled(doc, 'Estimated feature budget:', '18 hours (16% of total engagement)')
doc.add_paragraph()

# ─── Story 5.1 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 5.1: Written Summary Document', 2)

# Task 5.1.1
add_task_block(
    doc,
    task_id='Task 5.1.1',
    title='Assemble the Written Summary',
    hours='10h',
    what_text=(
        'Create the final client-facing document. Written for Jonas (strategic) and Julius '
        '(practitioner) simultaneously. 12-18 pages total.'
    ),
    how_items=[
        ('Document structure:', [
            'Executive Summary (1 page): headline finding, top 5 recommendations, immediate action needed.',
            'Section 1: Your Competitive Position -- where Goals sits vs market (Feature 1 findings).',
            'Section 2: Regional Pricing Assessment -- country risk map plus recommendation cards (Feature 2).',
            'Section 3: Pricing Scenarios -- what-if modelling showing revenue impact of adjustments (Story 1.2).',
            'Section 4: Risk Register -- all identified risks with mitigation (Story 3.2).',
            'Section 5: Recommendations -- top 5 specific changes (Story 3.3).',
            'Section 6: Pricing Framework -- reusable methodology for future SKUs (Feature 4, if completed).',
            'Section 7: World Cup Opportunities -- quick-hit tactical observations (Feature 4, if completed).',
            'Section 8: Next Steps -- Phase 2 positioning.',
            'Appendix: Supporting Data -- full comparison tables, country matrix, competitor data.',
        ]),
        'Style: British English. No em dashes. Direct, specific, numbers-backed. Written for Jonas (strategic) and Julius (practitioner) simultaneously.',
    ],
    done_when_text=(
        'Complete document draft. All sections populated with real findings. Internally '
        'consistent. Ready for internal review.'
    )
)

# ─── Story 5.2 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 5.2: Internal Review and Polish', 2)

# Task 5.2.1
add_task_block(
    doc,
    task_id='Task 5.2.1',
    title='Quality Gate -- Internal Review',
    hours='4h',
    what_text=(
        'Glen and Tom review before it goes to client.'
    ),
    how_items=[
        ('Glen reviews for:', [
            'Domain accuracy: do the recommendations make sense for this game?',
            'Tone: does it demonstrate expertise without being condescending?',
            'Completeness: does it answer Jonas\'s core questions?',
        ]),
        ('Tom reviews for:', [
            'Commercial positioning: does it set up Phase 2 naturally?',
            'NBI brand consistency.',
            'Any contractual/legal issues.',
        ]),
        'Address all feedback and produce final version.',
        'Send to Jonas/Julius 24 hours before the review call.',
    ],
    done_when_text=(
        'Final document approved by Glen and Tom. Sent to client. Call scheduled.'
    )
)

# ─── Story 5.3 ───────────────────────────────────────────────────────────────
add_heading(doc, 'Story 5.3: Remote Review Session', 2)

# Task 5.3.1
add_task_block(
    doc,
    task_id='Task 5.3.1',
    title='Conduct Client Review Call',
    hours='4h (prep + call + follow-up)',
    what_text=(
        '30-45 minute call walking Jonas/Julius through findings and recommendations.'
    ),
    how_items=[
        ('Agenda:', [
            '5 min: Recap scope and methodology.',
            '10 min: Top 5 recommendations (start with highest impact).',
            '10 min: Regional pricing highlights (red countries, what needs to change).',
            '5 min: Pricing framework overview.',
            '5 min: Next steps / Phase 2.',
            '10 min: Q&A.',
        ]),
        ('Prepared answers for likely questions:', [
            '"Should we delay launch?" -- No. These are configuration changes, not architecture.',
            '"What if platforms reject?" -- Built within their tier systems. Low risk if compliance verified.',
            '"How confident are you?" -- State honestly per finding. Data-backed = high. Scenario-based = moderate.',
            '"Can we change later?" -- Down only without backlash. Hence "launch high, patch down."',
        ]),
        'Follow-up within 24 hours: summary email, adjusted recommendations if needed, confirmed next steps.',
    ],
    done_when_text=(
        'Call completed. Follow-up sent. Phase 2 interest level documented.'
    )
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A: RISK REGISTER
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Appendix A: Risk Register', 1)

add_para(doc, (
    'All 10 pricing risks identified across the engagement. Ranked by severity x likelihood. '
    'Each risk includes description, severity, likelihood, markets affected, and mitigation.'
), space_after=8)

risks = [
    {
        'num': '1',
        'title': 'P2W Perception at High Price Points',
        'severity': 'Critical',
        'likelihood': 'Medium',
        'markets': 'All markets',
        'desc': (
            'Community data shows "pay to win" is the #1 cited quit reason. If premium packs '
            'feel like they buy wins, pricing is irrelevant because players leave regardless of '
            'price level. The 80/20 design mitigates this, but player perception does not '
            'always reflect reality.'
        ),
        'mitigation': (
            'Frame all HC purchases as time-saving or cosmetic, never as competitive advantage. '
            'Ensure store messaging emphasises "customise" not "dominate." Monitor community '
            'sentiment post-launch and respond within 48 hours if P2W narrative emerges.'
        ),
    },
    {
        'num': '2',
        'title': 'Post-Launch Price Increase Impossibility',
        'severity': 'High',
        'likelihood': 'Medium',
        'markets': 'All console markets',
        'desc': (
            'If current prices prove too low, raising them on console requires re-certification '
            '(slow), generates player backlash, and attracts press coverage. This is a platform '
            'constraint, not just a commercial one.'
        ),
        'mitigation': (
            'Launch at the upper end of the acceptable range. Use "introductory pricing" framing '
            'explicitly so future adjustments are expected rather than surprising.'
        ),
    },
    {
        'num': '3',
        'title': 'Over-Pricing in Price-Sensitive Markets',
        'severity': 'High',
        'likelihood': 'High',
        'markets': 'Turkey, Brazil, Poland',
        'desc': (
            'Large player bases with low purchasing power. Over-pricing causes low conversion '
            'and missed World Cup revenue. Poland is already confirmed at +25% above FX in '
            'Goals\' own matrix.'
        ),
        'mitigation': (
            'Implement country-specific adjustments per Story 2.3 recommendation cards. '
            'Address before platform submission.'
        ),
    },
    {
        'num': '4',
        'title': 'Platform Pricing Rejection',
        'severity': 'Critical',
        'likelihood': 'Low (if compliance verified)',
        'markets': 'PlayStation, Xbox territories',
        'desc': (
            'Sony/Xbox reject submitted pricing, causing delay past the May 14 launch window. '
            'Non-standard tier selections are the most common rejection reason.'
        ),
        'mitigation': (
            'Task 2.2.1 verification before submission. Submit at least 2 weeks early. '
            'Julius to confirm whether a test submission has already been run.'
        ),
    },
    {
        'num': '5',
        'title': 'Discount Curve Too Shallow for Launch',
        'severity': 'Medium',
        'likelihood': 'Medium',
        'markets': 'All markets (whale spending globally)',
        'desc': (
            'New game needs to incentivise initial spending. Goals\' 18% max discount may not '
            'drive enough whale commitment early when there is no established brand trust.'
        ),
        'mitigation': (
            'Consider steepening top-tier discount to 22-25% for first 90 days, then '
            'normalise. This is reversible (discounts can decrease post-launch).'
        ),
    },
    {
        'num': '6',
        'title': 'Missing Entry-Tier Conversion',
        'severity': 'Medium',
        'likelihood': 'Medium',
        'markets': 'All markets (first-time buyers)',
        'desc': (
            '$5.99 minimum may be too high for first-time buyers who need a low-risk trial '
            'purchase. No sub-$5 entry point exists in the current ladder.'
        ),
        'mitigation': (
            'First-purchase bonus (Story 1.3) addresses this without adding a micro-tier. '
            'A one-time $3.99 starter pack breaks the barrier while keeping the standard '
            'ladder intact.'
        ),
    },
    {
        'num': '7',
        'title': 'Originals/Celebrity Pack Pricing Gap',
        'severity': 'Low (not launch-critical)',
        'likelihood': 'N/A',
        'markets': 'Post-launch, all markets',
        'desc': (
            'Premium HC-only items (documented in the Goals Studio Monetisation Design Doc) have no pricing framework. '
            'When Originals and Celebrity packs launch post-Day 1, Goals needs a method to '
            'price them consistently.'
        ),
        'mitigation': (
            'Feature 4 framework covers this. Flag for Phase 2 detailed work. Not launch-blocking.'
        ),
    },
    {
        'num': '8',
        'title': 'Netherlands Regulatory Risk for Player Packs',
        'severity': 'High',
        'likelihood': 'Medium',
        'markets': 'Netherlands (EUR zone)',
        'desc': (
            'Player packs that grant athletes with meaningful attribute variance (even within '
            'the 80/20 model) may be classified as gambling-adjacent under Dutch Kansspelautoriteit '
            '(KSA) rulings. EA FC has been fined in the Netherlands for similar mechanics.'
        ),
        'mitigation': (
            'Legal review of pack mechanic before Netherlands launch. Consider offering '
            'cosmetic-only packs in the Netherlands, or ensure pack contents are predetermined '
            'and visible before purchase. Research current KSA enforcement stance on F2P '
            'sports games.'
        ),
    },
    {
        'num': '9',
        'title': 'South Korea Drop Rate Disclosure Requirement',
        'severity': 'High',
        'likelihood': 'Certain (if launching in South Korea)',
        'markets': 'South Korea',
        'desc': (
            'Korean law mandates disclosure of exact probabilities for all randomised item '
            'purchases. Non-compliance blocks the Korean launch entirely.'
        ),
        'mitigation': (
            'Confirm drop rate disclosure is built into the pack UI before Korean submission. '
            'Goals\' player tier distribution (in player_pricing data) provides the rates -- '
            'ensure they are displayed in-game.'
        ),
    },
    {
        'num': '10',
        'title': 'Non-Seasonal Demand Sustainability',
        'severity': 'Medium (long-term)',
        'likelihood': 'High',
        'markets': 'All markets (long-term)',
        'desc': (
            'Without seasonal resets or a battle pass, there is no guaranteed per-player spend '
            'floor per quarter. Player ageing is a gradual sink, leading to flat rather than '
            'pulsing revenue. This is the fundamental trade-off of the anti-P2W positioning.'
        ),
        'mitigation': (
            'Plan regular cosmetic refresh cadence (monthly kit/celebration drops). Leverage '
            'real-world football events as demand drivers. Use mode variety to drive squad '
            'depth demand. This is Phase 2 work but should be acknowledged in the deliverable.'
        ),
    },
]

for risk in risks:
    h = doc.add_heading(f'Risk {risk["num"]}: {risk["title"]}', level=2)

    add_table(doc,
        ['Field', 'Detail'],
        [
            ['Severity', risk['severity']],
            ['Likelihood', risk['likelihood']],
            ['Markets Affected', risk['markets']],
        ]
    )

    add_labelled(doc, 'Description:', risk['desc'])
    add_labelled(doc, 'Mitigation:', risk['mitigation'])
    doc.add_paragraph()

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B: EXECUTION TIMELINE (Option A: AI-Assisted)
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Appendix B: Execution Timeline -- Option A (AI-Assisted, 5 Calendar Days)', 1)

add_para(doc, (
    'Option A assumes AI handles the bulk of the research workload (competitive store '
    'scraping, country categorisation, data gathering). The team focuses on synthesis, '
    'judgement, and deliverable writing. This achieves delivery within 5 calendar days plus '
    'the review call.'
), space_after=8)

add_table(doc,
    ['Day', 'Tasks', 'Estimated Hours', 'AI Support'],
    [
        ['Day 1', 'T1.0.1, T1.0.2, T1.1.1 start', '~8h', 'T1.1.1 data gathering, T2.1.1 country categorisation'],
        ['Day 2', 'T1.1.2, T1.1.3, T1.1.4, T1.2.1 start, T2.1.2 start', '~20h', 'T1.1.4 calculations, T2.1.2 store scraping'],
        ['Day 3', 'T1.2.1 finish, T1.3.1, T2.1.2 finish, T2.2.1', '~20h', 'T2.3.1 card drafts for review'],
        ['Day 4', 'T3.1.1, T3.2.1, T3.3.1, T4.1.1, T4.2.1, T4.3.1', '~19h', 'T4.1.1, T4.2.1 drafts'],
        ['Day 5', 'T5.1.1, T5.2.1', '~14h', '--'],
        ['Day 6', 'T5.3.1 review call + follow-up', '~4h', '--'],
    ]
)

add_para(doc, (
    'Note on Option B (Full Human, 8-10 Calendar Days): Without AI research support, the '
    'team works at realistic pace and delivery slips to approximately 6 May. This may conflict '
    'with the platform submission deadline. A decision between Option A and Option B is '
    'required before Day 1 commences.'
), space_after=8)

add_heading(doc, 'Dependencies and Blockers', 2)

add_table(doc,
    ['Dependency', 'Status', 'Fallback'],
    [
        ['Goals\' Pricing Matrix data', 'In hand', 'N/A'],
        ['HC-to-item value mapping', 'Derive from docs (Task 1.0.1)', 'Ask Julius to confirm if unclear'],
        ['Which discount curve is final', 'BLOCKER -- UNKNOWN', 'Cannot proceed with T1.1.2 until confirmed'],
        ['Platform submission deadline', 'BLOCKER -- UNKNOWN', 'Cannot plan timeline until confirmed'],
        ['Day 1 store contents', 'BLOCKER -- needed for T1.0.1', 'Ask Julius'],
        ['Competitor regional pricing', 'Research needed (Task 2.1.2)', 'SteamDB + PS Store web access'],
        ['Platform tier validation', 'Partially in hand (matrix has PS wholesale)', 'Ask Julius to confirm valid tiers'],
        ['Session telemetry from Julius', 'Requested, pending', 'Proceed without -- note limitation in deliverable'],
        ['Call scheduling with Jonas', 'Not confirmed', 'Send availability request Day 1'],
        ['Published conversion/ARPPU data', 'Research needed (Task 1.0.2)', 'EA/Konami quarterly reports, Sensor Tower'],
    ]
)

add_heading(doc, 'Quality Gates', 2)

add_numbered_list(doc, [
    'Before Day 1 starts: Julius confirms discount curve choice, submission deadline, and Day 1 store contents. Without these, Tasks 1.0.1 and 1.1.2 cannot start properly.',
    'End of Day 1: Can we answer "what does each HC tier buy?" If not, escalate to Julius immediately.',
    'End of Day 2: Competitive position map draft complete. Do we know where Goals sits vs market?',
    'End of Day 3 (CRITICAL): Glen reviews draft top 5 recommendations. Go/no-go on direction before synthesis.',
    'End of Day 4: Full deliverable content complete (pending assembly). Feature 4 status: complete or cut.',
    'Day 5: Final document to Glen/Tom for review. Send to client by end of Day 5 or first thing Day 6.',
])

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
